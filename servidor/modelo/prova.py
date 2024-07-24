from docx import Document
from docx.oxml.table import CT_TcPr
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import random
import os
import base64
import qrcode
from setup.env import env

LETRAS = "abcde"
env = env()

def criar_run(texto, fonte, bold):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run(texto)
        run.font.name = fonte
        run.font.bold = bold
        para._element.get_or_add_pPr().spacing_after = Pt(0)      
        para._element.get_or_add_pPr().spacing_before = Pt(0)
    
        return run._element
    
def criar_sessao(num_colunas, cabecalho=False):
        sessao = OxmlElement("w:sectPr")
        
        layout = OxmlElement("w:type")
        layout.set(qn("w:val"), "continuous") 
        sessao.append(layout)
        
        colunas = OxmlElement("w:cols")
        colunas.set(qn("w:space"), "100")  
        colunas.set(qn("w:num"), str(num_colunas))
        sessao.append(colunas)
        if cabecalho:
            margem = OxmlElement("w:pgMar")
            margem.set(qn("w:top"), "720")  
            margem.set(qn("w:right"), "320")  
            margem.set(qn("w:bottom"), "320")  
            margem.set(qn("w:left"), "320")  
            margem.set(qn("w:header"), "720")  
            margem.set(qn("w:footer"), "320")  
            margem.set(qn("w:gutter"), "0")  
            sessao.append(margem)
        else:
            margem = OxmlElement("w:pgMar")
            margem.set(qn("w:top"), "720")  
            margem.set(qn("w:right"), "720")  
            margem.set(qn("w:bottom"), "720")  
            margem.set(qn("w:left"), "720")  
            margem.set(qn("w:header"), "720")  
            margem.set(qn("w:footer"), "720")  
            margem.set(qn("w:gutter"), "0")  
            sessao.append(margem)
        tamanho = OxmlElement("w:pgSz")
        tamanho.set(qn("w:w"), "12240")  
        tamanho.set(qn("w:h"), "15840")  
        sessao.append(tamanho)
        
        return sessao   
    
def criar_cabecalho(cabecalho, dados, aluno, id):
    mapa = {
        '<nome>': aluno["nome"],
        '<matricula>': aluno["matricula"],
        '<turma>': aluno["turma"],
        '<titulo>': f"Avaliação {dados["tipo"].capitalize()} de {dados["disciplina"]} do {dados["bimestre"]}º Bimestre"
    }
    for linha in cabecalho.rows:
        for celulas in linha.cells:
            for paragrafo in celulas.paragraphs:
                for marcacao, valor in mapa.items():
                    if marcacao in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(marcacao, str(valor))
                        for run in paragrafo.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,  
        box_size=10,  
        border=0, 
    )
    qr.add_data(id + ";" + aluno["matricula"])
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    img.save(env.STATIC + "qr.png")
    
    celulas_com_imagens = [
        (0, 0, 1.4,  'univap.png'),
        (0, 12, 2.4, 'fve.png'),
        (6, 11, 9.6, 'assinatura.png'),
        (7, 1, 15.2, 'observacoes.png'),
        (5, 0, 1.7,  'qr.png'),
        (3, 1, 0.33, 'letras.png'),
        (3, 2, 0.33, 'letras.png'),
        (3, 3, 0.33, 'letras.png'),
        (3, 4, 0.33, 'letras.png'),
        (3, 5, 0.33, 'letras.png'),
        (3, 6, 0.33, 'letras.png'),
        (3, 7, 0.33, 'letras.png'),
        (3, 8, 0.33, 'letras.png'),
        (3, 9, 0.33, 'letras.png'),
        (3, 10, 0.33,'letras.png')
    ]
    
    for linha, coluna, comprimento, imagem in celulas_com_imagens:
        if linha < len(cabecalho.rows) and coluna < len(cabecalho.rows[linha].cells):
            rota = env.STATIC + imagem
            cell = cabecalho.cell(linha, coluna)
        

            comprimento = Cm(comprimento)
          
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = paragraph.add_run()
            run.font.name = 'Arial'
            run.font.size = Pt(2)
            run.add_picture(rota, width=comprimento)
            
    return cabecalho

def embaralhar_alternativas(lista, posicao):
    item = lista[posicao]
    random.shuffle(lista)
    nova_posicao = lista.index(item)
    return nova_posicao

def embaralhar_perguntas(perguntas):
    for pergunta in perguntas:
        pergunta["indice_resposta"] = embaralhar_alternativas(pergunta["alternativas"], pergunta["indice_resposta"])
    random.shuffle(perguntas)        
    
def ajustar_pdf(caminho):
        reader = PdfReader(caminho)
        if len(reader.pages) % 2 != 0:
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            writer.add_blank_page()
            with open(caminho, 'wb') as output_pdf:
                writer.write(output_pdf)
                           
def Prova(arquivo, dados, aluno, colunas, perguntas, word, id):
    documento = Document(arquivo)    
    corpo = documento.element.body
    corpo.clear()
    
    p = documento.add_paragraph()._element
    propriedades = OxmlElement("w:pPr")
    propriedades.append(criar_sessao(1, True))
    p.insert(0, propriedades)

    corpo.insert(0, Document(env.STATIC + "cabecalho.docx").tables[0]._element)

    sessao = criar_sessao(colunas)
    
    embaralhar_perguntas(perguntas)
    gabarito = ""
    for num_pergunta, pergunta in enumerate(perguntas): 
        marcado = False
        for elemento in pergunta['enunciado']._element.getchildren():
            if not isinstance(elemento, CT_TcPr): 
                if isinstance(elemento, CT_P):
                    pPr = elemento.get_or_add_pPr()
                    pPr.spacing_before = Pt(6)
                    pPr.spacing_after = Pt(2)
                    
                    if not marcado:
                        numero = criar_run(f"{num_pergunta + 1}. ", "Arial", True)
                        elemento.insert(1, numero)
                        marcado = True
                sessao.append(elemento)
        
        for num_alternativa, alternativa in enumerate(pergunta['alternativas']):
            marcado = False
            for elemento in alternativa._element.getchildren():
                if not isinstance(elemento, CT_TcPr): 
                    if isinstance(elemento, CT_P):
                        pPr = elemento.get_or_add_pPr()
                        pPr.spacing_before = Pt(2)
                        pPr.spacing_after = Pt(2)
                        
                        if not marcado:
                            letra = criar_run(f"{LETRAS[num_alternativa]}) ", "Arial", True)
                            elemento.insert(1, letra)
                            marcado = True
                    sessao.append(elemento)
        gabarito += LETRAS[pergunta["indice_resposta"]] 
    corpo.append(sessao)
    criar_cabecalho(documento.tables[0], dados, aluno, id)
    rota = env.STATIC + "temp/temp"
    documento.save(rota + ".docx")
    doc = word.Documents.Open(os.path.realpath(rota + ".docx"))
    doc.SaveAs(os.path.realpath(rota + ".pdf"), FileFormat=17)
    doc.Close()
    ajustar_pdf(rota + ".pdf")
    with open(rota + ".pdf", "rb") as pdf_file:
        arquivo = base64.b64encode(pdf_file.read()).decode('utf-8')
        
    return {"matricula": aluno["matricula"], "documento": arquivo, "gabarito": gabarito}