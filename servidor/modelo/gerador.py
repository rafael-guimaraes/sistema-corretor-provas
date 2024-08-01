from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
from comtypes import client
from PyPDF2 import PdfMerger
import traceback
import os
import base64
from copy import deepcopy
from modelo.prova import Prova
from setup.env import env
from Emitter import emit
env = env()

from time import time

from io import BytesIO

class Gerador():
    def __init__(self,socket_connection, arquivo, dados, colunas, id) -> None:
        """
            arquivo (str) => Caminho para a rota do arquivo word da prova.
            alunos (list) => Lista contendo os alunos que realizaram a prova.
            colunas (int) => Numero de colunas divindo a prova.
        """
        self._socket_connection = socket_connection
        print(env.CHECKED+ __file__.replace(env.DIRECTORY,"") + " | Gerador.__init__ -> self._socket_connection", self._socket_connection)
        self.arquivo = arquivo
        self.dados = dados
        self.colunas = colunas
        self.id = id

    def ler_perguntas(self):
        T0 = time()
        perguntas = []
        
        doc = Document(self.arquivo)
        
        for tabela in doc.tables:
            for linha in tabela.rows:
                celula = linha.cells[0]
                enunciado = None
                alternativas = []
                for num_campo, campo in enumerate(celula.tables):
                    if num_campo == 0:
                        enunciado = campo.rows[0].cells[0]
                    else:
                        for alternativa in campo.rows:
                            alternativas.append(alternativa.cells[0])
                
                perguntas.append({
                    "enunciado": enunciado,
                    "alternativas": alternativas,
                    "indice_resposta": 0
                })
        return perguntas
       
   
    async def criar_provas(self, alunos, perguntas):
        if not alunos: return []
        provas_criadas = []
        
        word = client.CreateObject('Word.Application')
        num_alunos = len(alunos)
        await emit("updateProgress.createProva",self._socket_connection,{"total":num_alunos,"progress":0})
        for i,aluno in enumerate(alunos):
            T0 = time()
            prova =  Prova(self.arquivo, self.dados, aluno, self.colunas, deepcopy(perguntas), word, self.id)
            provas_criadas.append(prova)
            print(env.CHECKED+ __file__.replace(env.DIRECTORY,"") + " | Gerador.criar_provas() -> Objecto socket: " + str(self._socket_connection))
            print(env.SUCCESS+ __file__.replace(env.DIRECTORY,"") + " | Gerador.criar_provas() -> Time: " + str(time() - T0) + "s")
            await emit("updateProgress.createProva",self._socket_connection,{"total":num_alunos,"progress":i+1})

            
        word.quit()
        return provas_criadas
    
    def criar_exemplo(self, perguntas):
        word = client.CreateObject('Word.Application')
        aluno = {'nome': 'Fulano Braga da Silva Costa', 'matricula': '50220000', 'turma': '3H'} 
        T0 = time()
        prova =  Prova(self.arquivo, self.dados, aluno, self.colunas, deepcopy(perguntas), word, self.id)
        word.quit()
        print(env.CHECKED+ __file__.replace(env.DIRECTORY,"") + " | Gerador.criar_exemplo() -> Time: " + str(time() - T0) + "s")

        return [ prova ]
    
    def gerar_impressao(self, provas, saida_arquivo):
        T0 = time()
        merger = PdfMerger()
        
        for prova in provas:
            pdf_data = base64.b64decode(prova["documento"])
            
            pdf_file = BytesIO(pdf_data)
            
            merger.append(pdf_file)

        merger.write(saida_arquivo)
        merger.close()
        print(env.CHECKED+ __file__.replace(env.DIRECTORY,"") + " | Gerador.gerar_impressao() -> Time: " + str(time() - T0) + "s")

       
        
    def lista_de_chamada(self, alunos):
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        paragrafo = doc.add_paragraph('Lista de Chamada')
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fonte = paragrafo.style.font
        fonte.name = 'Arial'

        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = 'Table Grid'

        larguras = [0.5, 10, 3] 
        for i, largura in enumerate(larguras):
            tabela.columns[i].width = Inches(largura)
        cabecalho = tabela.rows[0].cells
        cabecalho[0].text = 'Matrícula'
        cabecalho[1].text = 'Nome'
        cabecalho[2].text = 'Assinatura'
        for aluno in alunos:
            linha = tabela.add_row().cells
            linha[0].text = aluno['matricula']
            linha[1].text = aluno['nome']
            linha[2].text = ''

        return doc
       